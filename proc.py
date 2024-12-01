import os
import docx
from docx import Document
from tqdm import tqdm
# from moviepy import concatenate_videoclips, VideoFileClip

from llm_mapreduce.utils import read_yaml
from llm_mapreduce.pipeline import BasePipeline

OFFLINE_FILES_PATH = "data/offline_files"
ONLINE_FILES_PATH = "data/online_files"
OUTPUT_PATH = "output"
MAPREDUCE_CONFIG = "llm_mapreduce/sum.yaml"
QUERY_PROMPT = """总结报告内容，要求内容不能是简单的课程笔记，而是能全面反映每次讲座的收获、心得体会及建议等。"""
CHUNK_SIZE = 4096


def preprocess():
    # Part 1: Read files from data/offline_files
    offline_files = []
    for file in os.listdir(OFFLINE_FILES_PATH):
        file_path = os.path.join(OFFLINE_FILES_PATH, file)
        if file.endswith('.docx'):
            doc = docx.Document(file_path)
            file_name = file[:-5]
            offline_files.append((file_name, "\n".join([p.text for p in doc.paragraphs])))
        elif file.endswith('.txt'):
            file_name = file[:-4]
            with open(file_path, 'r', encoding='utf-8') as f:
                offline_files.append((file_name, f.read()))
        else:
            raise NotImplementedError
    
    # Part 2: Process files and folders in data/online_files
    mp4_file_list = []
    # for item in os.listdir(ONLINE_FILES_PATH):
    #     item_path = os.path.join(ONLINE_FILES_PATH, item)
    #     if os.path.isdir(item_path):
    #         if os.path.exists(os.path.join(ONLINE_FILES_PATH, f"{item}.mp4")):
    #             continue
    #         mp4_clips = []
    #         for sub_item in sorted(os.listdir(item_path)):
    #             sub_item_path = os.path.join(item_path, sub_item)
    #             if sub_item.endswith('.mp4'):
    #                 mp4_clips.append(VideoFileClip(sub_item_path))
    #         if mp4_clips:
    #             # Concatenate and save as the folder name
    #             combined_clip = concatenate_videoclips(mp4_clips)
    #             output_path = os.path.join(ONLINE_FILES_PATH, f"{item}.mp4")
    #             combined_clip.write_videofile(output_path, codec="libx264")
    #             combined_clip.close()
    #             mp4_clips = [clip.close() for clip in mp4_clips]  # Close all clips after use
    #             mp4_file_list.append(f"{item}.mp4")
    #     elif item.endswith('.mp4'):
    #         mp4_file_list.append(item)

    return offline_files, mp4_file_list

def summarize(content, cid):
    map_reduce_config = read_yaml(MAPREDUCE_CONFIG)
    pipeline = BasePipeline(map_reduce_config, print_intermediate_path=os.path.join(OUTPUT_PATH, "log.txt"), doc_id=cid)
    result = pipeline.run(doc=content, question=QUERY_PROMPT, chunk_size=CHUNK_SIZE)
    return result.strip()

def save_paragraphs_to_docx(data):
    """
    Args:
        data (list of tuples): (title, content)。
    """
    filename = r"研究生素养课报告－姓名－学号－联系电话－Email地址.docx"
    os.makedirs(OUTPUT_PATH, exist_ok=True)
    
    doc = Document()
    for title, content in data:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    output_path = os.path.join(OUTPUT_PATH, filename)
    doc.save(output_path)
    
content_list, video_list = preprocess()
print([len(content[1]) for content in content_list])
content_list = [(content[0], summarize(content[1], cid)) for cid, content in tqdm(enumerate(content_list), desc="Summarizing Files")]
save_paragraphs_to_docx(content_list)
